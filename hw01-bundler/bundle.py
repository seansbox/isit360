import os
import sys
from glob import glob
from docx import Document
from docx.shared import Inches, Pt
from termcolor import colored

if getattr(sys, 'frozen', False):
    app_path = os.path.dirname(sys.executable)
elif __file__:
    app_path = os.path.dirname(__file__)
os.chdir(app_path)

types = [
    '**/*.md', '**/*.toml', '**/Pipfile',
    '**/*.py', '**/*.htm', '**/*.html',
    '**/*.csv', '**/*.json', '**/*.xml',
    '**/*.png', '**/*.jpg'
]

files = []
for t in types: files.extend(glob(t, recursive=True))
files = sorted(files)

document = Document()

print(colored(f"Bundling files...", 'white', attrs=['bold']))
for file in files:
    if 'build' in file or 'dist' in file or '__init__.py' in file: continue
    if 'manage.py' in file or 'asgi.py' in file or 'wsgi.py' in file: continue
    print(' - ' + colored(f"{file}", 'green'))
    document.add_heading(file, 0)
    if '.png' in file or '.jpg' in file:
        document.add_picture(file, width=Inches(6))
    else:
        with open(file) as f:
            document.add_paragraph(f.read()).style.font.size = Pt(9)
    document.add_page_break()

document.save('bundle.docx')